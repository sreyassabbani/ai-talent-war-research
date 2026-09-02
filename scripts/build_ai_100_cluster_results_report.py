"""Build the AI-100 clustering-results report from the published run snapshot."""

from __future__ import annotations

import csv
import json
import sys
from collections import Counter, defaultdict
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
SNAPSHOT = ROOT / "data" / "published" / "ai_100_overnight"
DOCS = ROOT / "docs"
MARKDOWN_OUT = DOCS / "ai_100_clustering_results_report.md"
DOCX_OUT = DOCS / "ai_100_clustering_results_report.docx"


TOPIC_LABELS = {
    "topic_5787d58e": "Topic 1: worker/productivity language",
    "topic_1a765e60": "Topic 2: workforce-management language",
    "topic_0036981e": "Topic 3: generic AI/company language",
}


def read_csv(name: str) -> list[dict[str, str]]:
    with (SNAPSHOT / name).open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def short(text: str, limit: int = 220) -> str:
    clean = " ".join(text.split())
    return clean if len(clean) <= limit else clean[: limit - 1].rstrip() + "…"


def load_data() -> dict[str, object]:
    manifest = read_csv("frozen_ai_manifest.csv")
    assignments = read_csv("topic_assignments.csv")
    deal_topics = read_csv("deal_by_topic.csv")
    topic_summary = read_csv("topic_summary.csv")
    diagnostics = json.loads((SNAPSHOT / "analysis_manifest.json").read_text(encoding="utf-8"))
    quality = json.loads((SNAPSHOT / "state.json").read_text(encoding="utf-8"))
    by_id = {row["deal_id"]: row for row in manifest}
    topic_counts = Counter(row["dominant_topic"] for row in assignments)
    topic_deals: dict[str, set[str]] = defaultdict(set)
    for row in assignments:
        topic_deals[row["dominant_topic"]].add(row["deal_id"])
    examples: dict[str, list[str]] = defaultdict(list)
    for row in sorted(assignments, key=lambda item: (-float(item["dominant_weight"]), item["passage_id"])):
        topic = row["dominant_topic"]
        if len(examples[topic]) < 3:
            examples[topic].append(short(row["supporting_excerpt"]))
    return {
        "manifest": manifest,
        "by_id": by_id,
        "assignments": assignments,
        "deal_topics": deal_topics,
        "topic_summary": topic_summary,
        "topic_counts": topic_counts,
        "topic_deals": topic_deals,
        "examples": examples,
        "diagnostics": diagnostics,
        "quality": quality,
    }


def deal_label(row: dict[str, str]) -> str:
    return f"{row['acquirer_name']} - {row['target_name']}"


def build_markdown(data: dict[str, object]) -> str:
    manifest = data["manifest"]
    by_id = data["by_id"]
    assignments = data["assignments"]
    deal_topics = data["deal_topics"]
    topic_summary = data["topic_summary"]
    topic_counts = data["topic_counts"]
    topic_deals = data["topic_deals"]
    examples = data["examples"]
    diagnostics = data["diagnostics"]
    quality = data["quality"]

    machine = [row for row in manifest if row["verification_status"] == "qualifying_machine_verified_pending_human_review"]
    rejected = [row for row in manifest if row["verification_status"] == "not_qualifying_no_primary_source_found"]
    deal_topic_fields = [field for field in deal_topics[0] if field.startswith("topic_")] if deal_topics else []

    lines: list[str] = [
        "# AI-100 Candidate Expansion and Clustering Results Report",
        "",
        "Status date: 2026-08-31",
        "",
        "## Executive summary",
        "",
        "This report explains what the 100-candidate expansion actually produced, which deals reached the unsupervised analysis, how the passages were grouped, and why the current groups cannot yet be treated as substantive employee-retention categories.",
        "",
        f"The run contains **{len(manifest)} candidate rows**: 100 selected candidates and 19 reserves. Of those, **{len(machine)} are machine-qualified and pending human review**, while **{len(rejected)} have no primary-source qualification**. Only **{len(deal_topics)} deals produced included passages**, yielding **{len(assignments)} passages** for exploratory clustering. There are currently **0 human-verified qualifying AI transactions** in this expansion. The topic result is rejected for release because one deal contributes 44.44% of passages, above the 35% concentration threshold.",
        "",
        "**Actual funnel:** 100 selected candidate slots → 35 machine-qualified pending human review → 13 deals with included passages → 72 passages → 3 exploratory topics. This is not a completed dataset of 100 verified deals.",
        "",
        "> The model grouped similar language. It did not determine that the deals used the same retention strategy, that the provisions worked, or that employees stayed.",
        "",
        "## 1. What the 100-candidate expansion actually produced",
        "",
        "| Run level | Count | Meaning |",
        "| --- | ---: | --- |",
        f"| Selected candidate slots | 100 | Screening target; not 100 completed or verified deals. |",
        f"| Reserve candidates | 19 | Backup rows included in the combined run. |",
        f"| Total candidate manifest rows | {len(manifest)} | Includes qualifying, rejected, and unresolved rows. |",
        f"| Machine-qualified rows | {len(machine)} | Pending human review. |",
        f"| No-primary-source rows | {len(rejected)} | Not counted as qualifying transactions. |",
        f"| Retrieved documents | {quality['stages']['report']['retrieved_documents']} | Includes 9 failed individual document retrievals. |",
        f"| Employee passages | {len(assignments)} | Screened passages, not validated provisions. |",
        f"| Deals represented in topic assignments | {len(deal_topics)} | Only deals with passages entered the model output. |",
        "",
        "The complete 119-row manifest is preserved in [`data/published/ai_100_overnight/frozen_ai_manifest.csv`](../data/published/ai_100_overnight/frozen_ai_manifest.csv). The 100 selected rows and 19 reserves are candidate records, not 119 verified deals. The 35 machine-qualified rows still require human review, and only 13 deals produced passages for the current model output.",
        "",
        "### Deals that produced passages",
        "",
        "These are the 13 deals represented in the current topic assignments. They are a processed subset of the candidate expansion, not a complete 100-deal results set. The remaining machine-qualified rows did not produce included passages in this run.",
        "",
        "| Deal | Passages | Dominant topic | Dominant share |",
        "| --- | ---: | --- | ---: |",
    ]
    for row in sorted(deal_topics, key=lambda item: (-int(item["passage_count"]), item["deal_id"])):
        label = deal_label(by_id[row["deal_id"]]) if row["deal_id"] in by_id else row["deal_id"]
        topic = row["dominant_topic"]
        lines.append(
            f"| {label} | {row['passage_count']} | {TOPIC_LABELS.get(topic, topic)} | {float(row[topic]):.1%} |"
        )

    lines += [
        "",
        "### Why most candidates did not reach the model",
        "",
        "The unresolved/rejected rows were preserved rather than silently dropped. These rows did not reach the current passage/model stage. The main missingness reasons were:",
        "",
        "- 57 acquirers could not be resolved on EDGAR.",
        "- 18 rows mentioned the target but lacked paragraph-local AI evidence.",
        "- 7 rows had no target mention in retrieved documents.",
        "- 2 rows had no documents retrieved for the deal.",
        "",
        "These are evidence-screening outcomes, not proof that the underlying transaction was not AI-related.",
        "",
        "## 2. How the unsupervised model worked",
        "",
        "The expansion used the following sequence:",
        "",
        "1. Retrieve target-linked SEC or approved official-source documents.",
        "2. Extract and screen employee-related or transaction-linked passages.",
        "3. Remove exact duplicate text and exclude common English or HTML-layout tokens.",
        "4. Convert passages into word and bigram TF-IDF vectors.",
        "5. Fit NMF models for K = 3, 4, 5, 6, and 7 using a fixed seed.",
        "6. Select the candidate K with the best deterministic half-sample stability.",
        "7. Assign each passage a dominant topic and a normalized weight across all topics.",
        "8. Aggregate passage weights back to deals for descriptive comparison.",
        "9. Compare NMF assignments with cosine/average agglomerative clustering and leave-one-deal-out term stability.",
        "",
        "The current selected configuration was K = 3, with half-sample stability `0.2464`. The model used scikit-learn-compatible TF-IDF/NMF logic in [`src/tag_edgar/topics100.py`](../src/tag_edgar/topics100.py). This was lexical topic exploration, not semantic understanding.",
        "",
        "## 3. What the three clusters contain",
        "",
        "| Topic | Passage count | Deal coverage | Top terms | Why passages were grouped |",
        "| --- | ---: | ---: | --- | --- |",
    ]
    for row in topic_summary:
        topic = f"topic_{row['topic_id']}" if not row["topic_id"].startswith("topic_") else row["topic_id"]
        # topic_summary stores numeric IDs; align them with the exported topic labels by row order.
        if row["topic_id"] == "0":
            topic = "topic_5787d58e"
        elif row["topic_id"] == "1":
            topic = "topic_1a765e60"
        else:
            topic = "topic_0036981e"
        terms = ", ".join(term.strip() for term in row["top_terms"].split(";") if term.strip())
        if topic == "topic_5787d58e":
            reason = "Repeated Zebra/frontline-worker/productivity wording. All 20 passages come from the Zebra–Fetch source, so this is primarily document-specific vocabulary."
        elif topic == "topic_1a765e60":
            reason = "Repeated workforce-management, Nucleus Research, and management-heading language. This looks like mixed product/marketing text rather than one employee-protection mechanism."
        else:
            reason = "Common acquisition-announcement language such as AI, founder, CEO, data, platform, team, and said. It is the broadest and most generic grouping."
        lines.append(
            f"| {TOPIC_LABELS[topic]} | {topic_counts[topic]} | {len(topic_deals[topic])} | {terms} | {reason} |"
        )

    lines += [
        "",
        "### Representative language behind the groupings",
        "",
    ]
    for topic in ["topic_5787d58e", "topic_1a765e60", "topic_0036981e"]:
        lines.append(f"**{TOPIC_LABELS[topic]}**")
        for example in examples[topic]:
            lines.append(f"- {example}")
        lines.append("")

    lines += [
        "The word patterns explain the assignments mechanically: passages containing similar words receive more weight in the same NMF component. They do not establish that the passages have the same legal meaning. For example, `worker`, `productivity`, and `frontline` can describe a buyer's product marketing rather than employee retention.",
        "",
        "## 4. Deal-level clustering view",
        "",
        "The table below shows the normalized average topic weights used for deal-level comparison. These are descriptive model weights, not probabilities that a deal used a particular retention strategy.",
        "",
        "| Deal | Passages | Topic 1 | Topic 2 | Topic 3 | Dominant topic |",
        "| --- | ---: | ---: | ---: | ---: | --- |",
    ]
    for row in sorted(deal_topics, key=lambda item: (-int(item["passage_count"]), item["deal_id"])):
        label = deal_label(by_id[row["deal_id"]]) if row["deal_id"] in by_id else row["deal_id"]
        lines.append(
            f"| {label} | {row['passage_count']} | {float(row['topic_5787d58e']):.1%} | {float(row['topic_1a765e60']):.1%} | {float(row['topic_0036981e']):.1%} | {TOPIC_LABELS.get(row['dominant_topic'], row['dominant_topic'])} |"
        )

    lines += [
        "",
        "The most important pattern is concentration: Zebra–Fetch contributes 32 of the 72 passages. It therefore dominates Topic 1 and contributes substantially to Topic 2. Most other deals are assigned primarily to Topic 3 because their shorter official announcements share generic AI/company language.",
        "",
        "## 5. Diagnostics and why the result is not release-ready",
        "",
        "| Diagnostic | Result | Interpretation |",
        "| --- | ---: | --- |",
        f"| Selected topic count | K = {diagnostics['topic_count']} | Best candidate in the configured K range. |",
        f"| Half-sample stability | {diagnostics['topics_config']['seed']} seed; 0.2464 | Low exploratory stability; not a validated taxonomy. |",
        f"| NMF/agglomerative ARI | {diagnostics['agglomerative_ari']:.4f} | Some assignment agreement, but not semantic validation. |",
        f"| Maximum deal passage share | {diagnostics['max_deal_passage_share']:.2%} | Exceeds the {diagnostics['max_deal_passage_share_threshold']:.0%} concentration threshold. |",
        f"| Human topic review | Blank | No human-approved labels. |",
        "",
        "The separate 10-deal pilot found an additional corpus-quality problem: included-passage relevance was 72.0% against a 90% gate, and missed relevant content among excluded candidates was 5.33% against a below-5% gate. That audit was performed on the pilot corpus, not the 100-deal output, but it is a warning that the passage screen must be repaired before interpreting topic differences substantively.",
        "",
        "## 6. What we can responsibly say",
        "",
        "- The model found recurring lexical patterns in a screened set of transaction-linked passages.",
        "- Topic 1 is heavily associated with one Zebra–Fetch source and worker/productivity wording.",
        "- Topic 2 mixes workforce-management and company-content language and has limited interpretive clarity.",
        "- Topic 3 is a broad generic AI/company-announcement cluster spanning the 13 deals with passages.",
        "- The output is useful for identifying passages that need human review and for designing a better corpus.",
        "",
        "## 7. What we cannot claim",
        "",
        "- The clusters are not validated employee-retention categories.",
        "- A dominant topic does not mean a company used that retention strategy.",
        "- A disclosed bonus, equity award, benefit, or service condition does not prove retention.",
        "- Zero passages do not prove that no employee arrangement existed.",
        "- The 100 selected candidate slots are not 100 completed or verified AI transactions.",
        "- The model does not measure company concern, employee motivation, or causal workforce outcomes.",
        "",
        "## 8. Recommended next step",
        "",
        "1. Repair the passage screen using the 150-row human relevance audit.",
        "2. Separate official deal announcements from product-marketing and unrelated company pages.",
        "3. Rebuild the passage corpus with one stable deal/document/passage provenance chain.",
        "4. Rerun the topic model and concentration/stability diagnostics.",
        "5. Have two real reviewers code representative passages before naming or releasing any cluster.",
        "6. Only after human validation, decide whether to create a codebook and later supervised classifier.",
        "",
        "## Source artifacts",
        "",
        "- [`frozen_ai_manifest.csv`](../data/published/ai_100_overnight/frozen_ai_manifest.csv) — all 119 candidate/reserve rows.",
        "- [`deal_source_register.csv`](../data/published/ai_100_overnight/deal_source_register.csv) — source and evidence register.",
        "- [`topic_assignments.csv`](../data/published/ai_100_overnight/topic_assignments.csv) — passage-level assignments and supporting excerpts.",
        "- [`deal_by_topic.csv`](../data/published/ai_100_overnight/deal_by_topic.csv) — deal-level normalized topic weights.",
        "- [`topic_summary.csv`](../data/published/ai_100_overnight/topic_summary.csv) — top terms and stability fields.",
        "- [`analysis_manifest.json`](../data/published/ai_100_overnight/analysis_manifest.json) — configuration, hashes, and status.",
        "- [`docs/unsupervised_models_progress.md`](unsupervised_models_progress.md) — complete history of the unsupervised-model work.",
        "",
    ]
    return "\n".join(lines)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(size)


def set_table_widths(table, widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], font_size: int = 8) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        set_cell_text(cell, header, bold=True, size=font_size)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8EEF5")
        cell._tc.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            set_cell_text(cell, value, size=font_size)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = 0


def build_docx(data: dict[str, object]) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    manifest = data["manifest"]
    by_id = data["by_id"]
    assignments = data["assignments"]
    deal_topics = data["deal_topics"]
    topic_summary = data["topic_summary"]
    topic_counts = data["topic_counts"]
    topic_deals = data["topic_deals"]
    examples = data["examples"]
    diagnostics = data["diagnostics"]
    quality = data["quality"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in [
        ("Heading 1", 16, "1F4D78", 14, 6),
        ("Heading 2", 12.5, "1F4D78", 10, 4),
        ("Heading 3", 11, "365F91", 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "TAG INTERNSHIP | AI-100 CLUSTERING RESULTS"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    footer = section.footer.paragraphs[0]
    footer.text = "Exploratory analysis - source-backed but not human-verified"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("AI-100 Candidate Expansion and Clustering Results Report")
    run.font.name = "Arial"
    run.font.size = Pt(25)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("Which candidates reached the model, how the passages grouped, and why the current result remains exploratory")
    sr.font.name = "Arial"
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(80, 80, 80)
    meta = doc.add_paragraph("Status date: 2026-08-31 | Run: ai_100_overnight | Candidate expansion snapshot")
    meta.runs[0].font.name = "Arial"
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    def heading(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def para(text: str, bold_prefix: str | None = None) -> None:
        p = doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            p.add_run(bold_prefix).bold = True
            p.add_run(text[len(bold_prefix) :])
        else:
            p.add_run(text)

    def bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)

    def number(text: str) -> None:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)

    heading("Executive summary")
    para("This report explains what the 100-candidate expansion actually produced, which deals reached the unsupervised analysis, how the passages were grouped, and why the current groups cannot yet be treated as substantive employee-retention categories.")
    para(f"The run contains {len(manifest)} candidate rows: 100 selected candidates and 19 reserves. Of those, 35 are machine-qualified and pending human review, while 84 have no primary-source qualification. Only {len(deal_topics)} deals produced included passages, yielding {len(assignments)} passages for exploratory clustering. There are currently 0 human-verified qualifying AI transactions in this expansion. The topic result is rejected for release because one deal contributes 44.44% of passages, above the 35% concentration threshold.")
    para("Actual funnel: 100 selected candidate slots -> 35 machine-qualified pending human review -> 13 deals with included passages -> 72 passages -> 3 exploratory topics. This is not a completed dataset of 100 verified deals.")
    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.keep_together = True
    callout_run = callout.add_run("The model grouped similar language. It did not determine that the deals used the same retention strategy, that the provisions worked, or that employees stayed.")
    callout_run.bold = True
    callout_run.font.name = "Arial"
    callout_run.font.size = Pt(10)
    callout_run.font.color.rgb = RGBColor(31, 77, 120)
    callout_pr = callout._p.get_or_add_pPr()
    callout_shading = OxmlElement("w:shd")
    callout_shading.set(qn("w:fill"), "E8EEF5")
    callout_pr.append(callout_shading)

    heading("1. What the 100-candidate expansion actually produced")
    add_table(doc, ["Run level", "Count", "Meaning"], [
        ["Selected candidate slots", "100", "Screening target; not 100 completed or verified deals."],
        ["Reserve candidates", "19", "Backup rows included in the combined run."],
        ["Total candidate manifest rows", str(len(manifest)), "Includes qualifying, rejected, and unresolved rows."],
        ["Machine-qualified rows", "35", "Pending human review."],
        ["No-primary-source rows", "84", "Not counted as qualifying transactions."],
        ["Retrieved documents", str(quality["stages"]["report"]["retrieved_documents"]), "Includes 9 failed individual document retrievals."],
        ["Employee passages", str(len(assignments)), "Screened passages, not validated provisions."],
        ["Deals represented in topic assignments", str(len(deal_topics)), "Only deals with passages entered the model output."],
    ], [2700, 1200, 5460], 8.5)
    para("The complete 119-row manifest is preserved in data/published/ai_100_overnight/frozen_ai_manifest.csv. The 100 selected rows and 19 reserves are candidate records, not 119 verified deals. The 35 machine-qualified rows still require human review, and only 13 deals produced passages for the current model output.")

    heading("Deals that produced passages", 2)
    para("These are the 13 deals represented in the current topic assignments. They are a processed subset of the candidate expansion, not a complete 100-deal results set. The remaining machine-qualified rows did not produce included passages in this run.")
    deal_rows = []
    for row in sorted(deal_topics, key=lambda item: (-int(item["passage_count"]), item["deal_id"])):
        label = deal_label(by_id[row["deal_id"]]) if row["deal_id"] in by_id else row["deal_id"]
        deal_rows.append([label, row["passage_count"], TOPIC_LABELS.get(row["dominant_topic"], row["dominant_topic"]), f"{float(row[row['dominant_topic']]):.1%}"])
    add_table(doc, ["Deal", "Passages", "Dominant topic", "Share"], deal_rows, [3200, 1000, 3900, 1260], 7.5)

    heading("Why most candidates did not reach the model", 2)
    para("The unresolved/rejected rows were preserved rather than silently dropped. These rows did not reach the current passage/model stage. The main missingness reasons were:")
    for item in ["57 acquirers could not be resolved on EDGAR.", "18 rows mentioned the target but lacked paragraph-local AI evidence.", "7 rows had no target mention in retrieved documents.", "2 rows had no documents retrieved for the deal."]:
        bullet(item)
    para("These are evidence-screening outcomes, not proof that the underlying transaction was not AI-related.")

    heading("2. How the unsupervised model worked")
    para("The expansion used the following sequence:")
    for item in [
        "Retrieve target-linked SEC or approved official-source documents.",
        "Extract and screen employee-related or transaction-linked passages.",
        "Remove exact duplicate text and exclude common English or HTML-layout tokens.",
        "Convert passages into word and bigram TF-IDF vectors.",
        "Fit NMF models for K = 3, 4, 5, 6, and 7 using a fixed seed.",
        "Select the candidate K with the best deterministic half-sample stability.",
        "Assign each passage a dominant topic and a normalized weight across all topics.",
        "Aggregate passage weights back to deals for descriptive comparison.",
        "Compare NMF assignments with cosine/average agglomerative clustering and leave-one-deal-out term stability.",
    ]:
        number(item)
    para("The selected configuration was K = 3, with half-sample stability 0.2464. This was lexical topic exploration, not semantic understanding.")

    heading("3. What the three clusters contain")
    topic_rows = []
    for row in topic_summary:
        topic = ["topic_5787d58e", "topic_1a765e60", "topic_0036981e"][int(row["topic_id"])]
        if topic == "topic_5787d58e":
            reason = "Repeated Zebra/frontline-worker/productivity wording. All 20 passages come from the Zebra-Fetch source, so this is primarily document-specific vocabulary."
        elif topic == "topic_1a765e60":
            reason = "Repeated workforce-management, Nucleus Research, and management-heading language. This looks like mixed product/marketing text rather than one employee-protection mechanism."
        else:
            reason = "Common acquisition-announcement language such as AI, founder, CEO, data, platform, team, and said. It is the broadest and most generic grouping."
        topic_rows.append([TOPIC_LABELS[topic], str(topic_counts[topic]), str(len(topic_deals[topic])), row["top_terms"].replace(";", ", "), reason])
    add_table(doc, ["Topic", "Passages", "Deals", "Top terms", "Why grouped"], topic_rows, [1800, 800, 700, 2200, 3860], 7.2)
    heading("Representative language behind the groupings", 2)
    for topic in ["topic_5787d58e", "topic_1a765e60", "topic_0036981e"]:
        p = doc.add_paragraph()
        p.add_run(TOPIC_LABELS[topic]).bold = True
        for example in examples[topic]:
            bullet(example)
    para("The word patterns explain the assignments mechanically: passages containing similar words receive more weight in the same NMF component. They do not establish that the passages have the same legal meaning. For example, worker, productivity, and frontline can describe a buyer's product marketing rather than employee retention.")

    heading("4. Deal-level clustering view")
    para("These normalized average topic weights are descriptive model weights, not probabilities that a deal used a particular retention strategy.")
    deal_matrix_rows = []
    for row in sorted(deal_topics, key=lambda item: (-int(item["passage_count"]), item["deal_id"])):
        label = deal_label(by_id[row["deal_id"]]) if row["deal_id"] in by_id else row["deal_id"]
        deal_matrix_rows.append([label, row["passage_count"], f"{float(row['topic_5787d58e']):.1%}", f"{float(row['topic_1a765e60']):.1%}", f"{float(row['topic_0036981e']):.1%}", TOPIC_LABELS.get(row["dominant_topic"], row["dominant_topic"])])
    add_table(doc, ["Deal", "Passages", "Topic 1", "Topic 2", "Topic 3", "Dominant"], deal_matrix_rows, [2750, 750, 800, 800, 800, 3460], 7.2)
    para("The most important pattern is concentration: Zebra-Fetch contributes 32 of the 72 passages. It therefore dominates Topic 1 and contributes substantially to Topic 2. Most other deals are assigned primarily to Topic 3 because their shorter official announcements share generic AI/company language.")

    heading("5. Diagnostics and why the result is not release-ready")
    add_table(doc, ["Diagnostic", "Result", "Interpretation"], [
        ["Selected topic count", "K = 3", "Best candidate in the configured K range."],
        ["Half-sample stability", "0.2464", "Low exploratory stability; not a validated taxonomy."],
        ["NMF/agglomerative ARI", f"{diagnostics['agglomerative_ari']:.4f}", "Some assignment agreement, but not semantic validation."],
        ["Maximum deal passage share", f"{diagnostics['max_deal_passage_share']:.2%}", "Exceeds the 35% concentration threshold."],
        ["Human topic review", "Blank", "No human-approved labels."],
    ], [2800, 1400, 4560], 8.2)
    para("The separate 10-deal pilot found an additional corpus-quality problem: included-passage relevance was 72.0% against a 90% gate, and missed relevant content among excluded candidates was 5.33% against a below-5% gate. That audit was performed on the pilot corpus, not the 100-deal output, but it is a warning that the passage screen must be repaired before interpreting topic differences substantively.")

    heading("6. What we can responsibly say")
    for item in [
        "The model found recurring lexical patterns in a screened set of transaction-linked passages.",
        "Topic 1 is heavily associated with one Zebra-Fetch source and worker/productivity wording.",
        "Topic 2 mixes workforce-management and company-content language and has limited interpretive clarity.",
        "Topic 3 is a broad generic AI/company-announcement cluster spanning the 13 deals with passages.",
        "The output is useful for identifying passages that need human review and for designing a better corpus.",
    ]:
        bullet(item)

    heading("7. What we cannot claim")
    for item in [
        "The clusters are not validated employee-retention categories.",
        "A dominant topic does not mean a company used that retention strategy.",
        "A disclosed bonus, equity award, benefit, or service condition does not prove retention.",
        "Zero passages do not prove that no employee arrangement existed.",
        "The 100 selected candidate slots are not 100 completed or verified AI transactions.",
        "The model does not measure company concern, employee motivation, or causal workforce outcomes.",
    ]:
        bullet(item)

    heading("8. Recommended next step")
    for item in [
        "Repair the passage screen using the 150-row human relevance audit.",
        "Separate official deal announcements from product-marketing and unrelated company pages.",
        "Rebuild the passage corpus with one stable deal/document/passage provenance chain.",
        "Rerun the topic model and concentration/stability diagnostics.",
        "Have two real reviewers code representative passages before naming or releasing any cluster.",
        "Only after human validation, decide whether to create a codebook and later supervised classifier.",
    ]:
        number(item)

    heading("Source artifacts")
    for item in [
        "frozen_ai_manifest.csv - all 119 candidate/reserve rows.",
        "deal_source_register.csv - source and evidence register.",
        "topic_assignments.csv - passage-level assignments and supporting excerpts.",
        "deal_by_topic.csv - deal-level normalized topic weights.",
        "topic_summary.csv - top terms and stability fields.",
        "analysis_manifest.json - configuration, hashes, and status.",
        "docs/unsupervised_models_progress.md - complete history of the unsupervised-model work.",
    ]:
        bullet(item)

    doc.core_properties.title = "AI-100 Candidate Expansion and Clustering Results Report"
    doc.core_properties.subject = "Deals, unsupervised topic assignments, and interpretation limits"
    doc.core_properties.author = "TAG Internship research project"
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def main() -> int:
    data = load_data()
    MARKDOWN_OUT.parent.mkdir(parents=True, exist_ok=True)
    MARKDOWN_OUT.write_text(build_markdown(data), encoding="utf-8")
    try:
        build_docx(data)
    except ModuleNotFoundError as exc:
        if exc.name == "docx":
            print("Markdown report written; python-docx is unavailable for DOCX creation.", file=sys.stderr)
            return 2
        raise
    print(f"Wrote {MARKDOWN_OUT}")
    print(f"Wrote {DOCX_OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
