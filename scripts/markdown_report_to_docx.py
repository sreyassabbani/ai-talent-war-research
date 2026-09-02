"""Render a generated Markdown report as a Word document.

The Markdown file is the source of truth; this only changes the container, so the DOCX cannot
carry a number the Markdown does not. Handles the subset of Markdown the report generator emits:
headings, paragraphs, bullet lists, block quotes, fenced code, and pipe tables.

Requires python-docx, which is optional tooling and not a package dependency:
    pip install python-docx

Usage:
    python scripts/markdown_report_to_docx.py docs/report.md docs/report.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_CODE = re.compile(r"`([^`]+)`")


def add_rich_text(paragraph: object, text: str) -> None:
    """Write text, rendering bold, inline code, and links as their visible text."""
    text = _LINK.sub(r"\1", text)
    position = 0
    for match in _BOLD.finditer(text):
        if match.start() > position:
            _add_plain(paragraph, text[position : match.start()])
        run = paragraph.add_run(_CODE.sub(r"\1", match.group(1)))  # type: ignore[attr-defined]
        run.bold = True
        position = match.end()
    if position < len(text):
        _add_plain(paragraph, text[position:])


def _add_plain(paragraph: object, text: str) -> None:
    for index, piece in enumerate(_CODE.split(text)):
        if not piece:
            continue
        run = paragraph.add_run(piece)  # type: ignore[attr-defined]
        if index % 2 == 1:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip("|").split("|")]


def add_table(document: DocumentObject, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.style = "Light Grid Accent 1"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for column in range(width):
            value = row[column] if column < len(row) else ""
            cell = cells[column]
            cell.text = ""
            add_rich_text(cell.paragraphs[0], value)
            if index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def convert(markdown_path: Path, docx_path: Path) -> None:
    document = Document()
    # python-docx types styles as BaseStyle, which has no font attribute; the Normal paragraph
    # style does. Set it defensively so a future python-docx cannot fail the whole conversion.
    style = document.styles["Normal"]
    font = getattr(style, "font", None)
    if font is not None:
        font.name = "Calibri"
        font.size = Pt(11)

    lines = markdown_path.read_text(encoding="utf-8").split("\n")
    index = 0
    pending: list[list[str]] = []

    def flush_table() -> None:
        nonlocal pending
        if pending:
            add_table(document, pending)
            document.add_paragraph()
            pending = []

    while index < len(lines):
        line = lines[index].rstrip()

        if is_table_row(line):
            cells = split_row(line)
            # The |---|---| separator carries no content.
            if not all(set(cell) <= set("-: ") and cell for cell in cells):
                pending.append(cells)
            index += 1
            continue
        flush_table()

        if line.startswith("```"):
            index += 1
            block: list[str] = []
            while index < len(lines) and not lines[index].startswith("```"):
                block.append(lines[index])
                index += 1
            index += 1
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            document.add_heading(line[level:].strip(), level=min(level, 4))
        elif line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(24)
            run = paragraph.add_run(_LINK.sub(r"\1", line[2:]))
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        elif line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, line[2:])
        elif re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            add_rich_text(paragraph, re.sub(r"^\d+\. ", "", line))
        elif line.strip() == "---":
            document.add_page_break()
        elif line.strip():
            # Consecutive non-empty lines are one wrapped paragraph in the source.
            buffer = [line.strip()]
            while (
                index + 1 < len(lines)
                and lines[index + 1].strip()
                and not lines[index + 1].startswith(("#", "|", ">", "- ", "* ", "```"))
                and not re.match(r"^\d+\. ", lines[index + 1])
            ):
                index += 1
                buffer.append(lines[index].strip())
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_rich_text(paragraph, " ".join(buffer))
        index += 1

    flush_table()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print((__doc__ or "").strip(), file=sys.stderr)
        return 2
    source, target = Path(argv[1]), Path(argv[2])
    if not source.exists():
        print(f"{source} not found.", file=sys.stderr)
        return 1
    convert(source, target)
    print(f"Wrote {target}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
